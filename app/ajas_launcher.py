#!/usr/bin/env python3
"""AJAS - AI-Powered Job Application System"""

import sys, os, json, shutil, platform, subprocess
import webbrowser, threading, time, logging
from pathlib import Path
from datetime import datetime, date

import re
import unicodedata

try:
    import requests as http_requests
except ImportError:
    http_requests = None

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError:
    openpyxl = None

try:
    from flask import (Flask, render_template, request,
                       jsonify, redirect, send_from_directory)
except ImportError:
    subprocess.call([sys.executable, "-m", "pip", "install",
                     "flask", "--quiet"])
    from flask import (Flask, render_template, request,
                       jsonify, redirect, send_from_directory)

APP_FOLDER  = "AI-Powered Job Application Files"
CONFIG_FILE = "ajas_config.json"
CLAUDE_URL  = "https://claude.ai"
PORT        = 7842

# Config is always stored in a fixed location in the user's home directory.
# This is separate from the AJAS files folder so we can always find it
# regardless of where the user chose to save their files.
CONFIG_PATH = Path.home() / ".ajas_config.json"

SUBFOLDERS = [
    "My CV/Tailored CVs",
    "Cover Letters",
    "Job Listings",
    "Application Tracker",
    "Thank You Emails",
]

app = Flask(__name__, template_folder="templates", static_folder="static")
app.secret_key = "ajas-local-only-2024"
app.config["SESSION_COOKIE_SECURE"] = False
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

# ── Helpers ───────────────────────────────────────────────────────────────────

def default_dir():
    return Path.home() / APP_FOLDER

def find_config():
    """Load config from fixed location."""
    if CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH) as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_config(cfg):
    cfg["last_updated"] = datetime.now().isoformat()
    with open(CONFIG_PATH, "w") as f:
        json.dump(cfg, f, indent=2)
    # Also ensure the AJAS files folder exists
    app_dir = cfg.get("app_dir")
    if app_dir:
        Path(app_dir).mkdir(parents=True, exist_ok=True)

def make_folders(path):
    for sub in SUBFOLDERS:
        (Path(path) / sub).mkdir(parents=True, exist_ok=True)

def open_in_finder(path):
    try:
        s = platform.system()
        cmd = ["open"] if s == "Darwin" else \
              ["explorer"] if s == "Windows" else ["xdg-open"]
        subprocess.Popen(cmd + [str(path)],
                         stdout=subprocess.DEVNULL,
                         stderr=subprocess.DEVNULL)
    except Exception:
        pass

def find_cv(folder):
    for ext in ["*.docx", "*.pdf", "*.doc"]:
        hits = list(Path(folder).glob(ext))
        if hits:
            return hits[0]
    return None

def missing_deps():
    missing = []
    for pkg in ["openpyxl", "docx", "requests"]:
        try:
            __import__(pkg)
        except ImportError:
            missing.append(pkg)
    return missing

def build_instructions(cfg, app_dir=None):
    """Build instructions text. Returns string (no longer writes to disk)."""
    lines = [
        "# AJAS - AI-Powered Job Application System",
        "## Claude Project Instructions",
        "## Generated: " + datetime.now().strftime("%Y-%m-%d %H:%M"),
        "",
        "---",
        "## WHO YOU ARE",
        "",
        "You are AJAS, an expert AI job application assistant. You guide the user",
        "through every stage of their job search: CV review, job discovery,",
        "CV tailoring, cover letter writing, application guidance, and tracking.",
        "You are professional, encouraging, honest, and efficient.",
        "",
        "---",
        "## USER PROFILE",
        "",
        "- CV File     : " + cfg.get("cv_filename", "not set"),
        "- Target Role : " + cfg.get("role_title",  "not specified"),
        "- Location    : " + cfg.get("location",     "not specified"),
        "- Salary Range: " + cfg.get("salary_range", "not specified"),
        "- Work Type   : " + cfg.get("work_type",    "Any"),
        "- Files at    : " + str(app_dir or cfg.get("app_dir", "")),
        "",
        "---",
        "## HOW TO START",
        "",
        "1. Greet the user warmly as AJAS",
        "2. Ask them to upload their CV or paste its contents",
        "3. Once received, run CV Review immediately (Module 1)",
        "4. After review, collect job preferences (Module 2)",
        "",
        "---",
        "## MODULE 1 - CV REVIEW",
        "",
        "Score the CV out of 100 across five areas (20 pts each):",
        "- Formatting and structure",
        "- ATS compatibility (keywords, standard headings)",
        "- Impact statements (achievements with numbers)",
        "- Length and relevance",
        "- Skills section completeness",
        "",
        "Provide: overall score, top 3 strengths, top 3 improvements,",
        "and a fully rewritten improved version.",
        "",
        "Ask: Would you like to use this improved version as your master CV,",
        "or would you prefer to update it yourself and come back?",
        "",
        "---",
        "## MODULE 2 - JOB PREFERENCES",
        "",
        "Collect: Target Role Title, Location, Salary Range, Work Type.",
        "Confirm back and say: Your preferences are saved.",
        "When ready to search, say search for jobs.",
        "",
        "---",
        "## MODULE 3 - JOB SEARCH",
        "",
        "Confirm preferences, then tell the user the AJAS app will",
        "run the search and update their job list in Excel.",
        "Once they have results, ask them to paste the job description",
        "for the role they want to apply for.",
        "",
        "---",
        "## MODULE 4 - CV TAILORING AND COVER LETTER",
        "",
        "1. Ask user to paste the job description",
        "2. Extract top 10 keywords and requirements",
        "3. Tailor master CV to match - preserve all facts, never invent",
        "4. Score tailored CV against the job (target 80+/100)",
        "5. If below 80, improve and re-score before presenting",
        "6. Write a fresh cover letter for this specific job (max 1 page)",
        "7. Present both for user approval",
        "8. Once approved, end your response with this EXACT text:",
        "   ---",
        "   Copy the TAILORED CV text above and the COVER LETTER text above,",
        "   then go back to AJAS and paste each into the boxes in Step 2.",
        "   ---",
        "",
        "---",
        "## MODULE 5 - APPLICATION GUIDANCE",
        "",
        "Guide the user field by field through the application form.",
        "Suggest answers for open-ended questions.",
        "Remind them to review before submitting.",
        "User always submits themselves - never auto-submit.",
        "",
        "---",
        "## MODULE 6 - APPLICATION TRACKING",
        "",
        "Tracker columns:",
        "Date Applied | Company | Role | Location | Work Type |",
        "Salary | Application URL | Status | Notes",
        "",
        "Status: Applied / Screening / Interview / Assessment /",
        "Final Round / Offer / Rejected / Withdrawn",
        "",
        "---",
        "## RULES",
        "",
        "- Never submit on the users behalf",
        "- Always get approval before finalising CV or cover letter",
        "- Never invent experience or qualifications",
        "- Be honest about CV weaknesses",
        "- Keep every response focused on the next action",
    ]
    return "\n".join(lines)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/favicon.ico")
def favicon():
    return send_from_directory(
        os.path.join(app.root_path, "static"),
        "favicon.ico", mimetype="image/vnd.microsoft.icon")

@app.route("/")
def index():
    cfg = find_config()
    if cfg.get("setup_complete"):
        return redirect("/dashboard")
    return redirect("/check")

# -- Requirements check -------------------------------------------------------

@app.route("/check")
def check():
    ver    = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
    py_ok  = sys.version_info >= (3, 9)
    miss   = missing_deps()
    return render_template("check.html",
        ver=ver, py_ok=py_ok, missing=miss, deps_ok=not miss)

@app.route("/check/install", methods=["POST"])
def install():
    miss = missing_deps()
    if miss:
        pm = {"docx": "python-docx"}
        names = [pm.get(p, p) for p in miss]
        subprocess.call([sys.executable, "-m", "pip",
                         "install", "--quiet"] + names)
    return redirect("/welcome")

# -- Welcome ------------------------------------------------------------------

@app.route("/welcome")
def welcome():
    return render_template("welcome.html")

# -- Folder setup -------------------------------------------------------------

@app.route("/folder")
def folder():
    return render_template("folder.html", default=str(default_dir()))

@app.route("/folder/browse", methods=["POST"])
def folder_browse():
    """Use a native OS folder picker and return the chosen path as JSON."""
    system = platform.system()
    chosen = None
    try:
        if system == "Darwin":
            result = subprocess.run([
                "osascript",
                "-e",
                'tell app "Finder" to activate\n'
                'tell app "Finder"\n'
                'set f to choose folder with prompt "Choose where to save your AJAS files:"\n'
                'return POSIX path of f\n'
                'end tell'
            ], capture_output=True, text=True, timeout=60)
            chosen = result.stdout.strip().rstrip("/")
        elif system == "Windows":
            # Use a simpler faster VBScript approach
            vbs = (
                'Set objShell = CreateObject("Shell.Application")\n'
                'Set objFolder = objShell.BrowseForFolder(0, '
                '"Choose where to save your AJAS files", 0)\n'
                'If Not objFolder Is Nothing Then\n'
                '    WScript.Echo objFolder.Self.Path\n'
                'End If'
            )
            import tempfile
            vbs_file = tempfile.NamedTemporaryFile(
                suffix=".vbs", mode="w", delete=False)
            vbs_file.write(vbs)
            vbs_file.close()
            try:
                result = subprocess.run(
                    ["cscript", "//nologo", vbs_file.name],
                    capture_output=True, text=True, timeout=30)
                chosen = result.stdout.strip()
            finally:
                try:
                    os.unlink(vbs_file.name)
                except Exception:
                    pass
    except Exception:
        pass

    if chosen:
        full = str(Path(chosen) / APP_FOLDER)
        return jsonify({"path": full})
    return jsonify({"path": ""})

@app.route("/folder/confirm", methods=["POST"])
def folder_confirm():
    path = request.form.get("folder_path", "").strip()
    if not path:
        path = str(default_dir())
    try:
        make_folders(path)
        cfg = find_config()
        cfg["app_dir"] = path
        save_config(cfg)
        return redirect("/cv")
    except Exception as e:
        return render_template("folder.html",
            default=str(default_dir()), error=str(e))

# -- CV upload ----------------------------------------------------------------

@app.route("/cv")
def cv():
    cfg     = find_config()
    app_dir = cfg.get("app_dir", str(default_dir()))
    cv_dir  = Path(app_dir) / "My CV"
    cv_file = find_cv(cv_dir)
    if cv_file:
        cfg["cv_filename"] = cv_file.name
        cfg["cv_path"]     = str(cv_file)
        save_config(cfg)
    return render_template("cv.html",
        cv=cv_file.name if cv_file else None,
        cv_dir=str(cv_dir))

@app.route("/cv/upload", methods=["POST"])
def cv_upload():
    cfg     = find_config()
    app_dir = cfg.get("app_dir", str(default_dir()))
    cv_dir  = Path(app_dir) / "My CV"
    cv_dir.mkdir(parents=True, exist_ok=True)

    f = request.files.get("cv_file")
    if not f or not f.filename:
        return render_template("cv.html", cv=None,
            cv_dir=str(cv_dir), error="No file selected.")

    # Only accept docx and pdf
    ext = Path(f.filename).suffix.lower()
    if ext not in [".docx", ".pdf", ".doc"]:
        return render_template("cv.html", cv=None,
            cv_dir=str(cv_dir),
            error="Please upload a .docx or .pdf file.")

    dest = cv_dir / f.filename
    f.save(str(dest))

    cfg["cv_filename"] = f.filename
    cfg["cv_path"]     = str(dest)
    save_config(cfg)
    return redirect("/prefs")

@app.route("/cv/open_folder")
def cv_open_folder():
    cfg = find_config()
    open_in_finder(Path(cfg.get("app_dir", str(default_dir()))) / "My CV")
    return redirect("/cv")

@app.route("/cv/refresh")
def cv_refresh():
    return redirect("/cv")

# -- Preferences --------------------------------------------------------------

@app.route("/prefs", methods=["GET", "POST"])
def prefs():
    cfg = find_config()
    if request.method == "POST":
        role   = request.form.get("role_title",   "").strip()
        loc    = request.form.get("location",      "").strip()
        salary = request.form.get("salary_range",  "").strip()
        wtype  = request.form.get("work_type",     "Any").strip()
        if not role or not loc or not salary:
            return render_template("prefs.html", cfg=cfg,
                error="Please fill in all three fields.")
        cfg.update({"role_title": role, "location": loc,
                    "salary_range": salary, "work_type": wtype})
        app_dir = cfg.get("app_dir", str(default_dir()))
        instr   = build_instructions(cfg, app_dir)
        already_setup = cfg.get("setup_complete", False)
        cfg["setup_complete"] = True
        save_config(cfg)
        if already_setup:
            return redirect("/prefs/saved")
        return redirect("/connect")
    return render_template("prefs.html", cfg=cfg, error=None)

# -- Claude connect -----------------------------------------------------------

@app.route("/connect", methods=["GET", "POST"])
def connect():
    cfg = find_config()
    if request.method == "POST":
        url = request.form.get("project_url", "").strip()
        if not url or not url.startswith("http"):
            return render_template("connect.html", cfg=cfg,
                error="Please paste your Claude Project URL — "
                      "this is required so AJAS can open it directly.")
        cfg["claude_project_url"] = url
        save_config(cfg)
        return redirect("/dashboard")
    return render_template("connect.html", cfg=cfg, error=None)

@app.route("/connect/open_instructions")
def open_instructions():
    cfg = find_config()
    p   = cfg.get("instructions_path", "")
    if p:
        open_in_finder(Path(p).parent)
    return redirect("/connect")

# -- Dashboard ----------------------------------------------------------------

@app.route("/dashboard")
def dashboard():
    return render_template("dashboard.html", cfg=find_config())

@app.route("/dashboard/open_claude")
def open_claude():
    cfg = find_config()
    url = cfg.get("claude_project_url", CLAUDE_URL)
    # Open reliably on all platforms
    try:
        if platform.system() == "Windows":
            subprocess.Popen(
                ["powershell", "-Command", f"Start-Process '{url}'"],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", url],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        else:
            subprocess.Popen(["xdg-open", url],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception:
        webbrowser.open(url)
    return redirect("/dashboard")

@app.route("/dashboard/open_folder")
def open_folder_route():
    cfg = find_config()
    open_in_finder(cfg.get("app_dir", str(default_dir())))
    return redirect("/dashboard")

@app.route("/dashboard/tracker")
def dashboard_tracker():
    return redirect("/tracker")

@app.route("/dashboard/update_prefs")
def update_prefs():
    return redirect("/prefs")

@app.route("/dashboard/update_cv")
def update_cv():
    cfg = find_config()
    cfg.pop("cv_path", None)
    cfg.pop("cv_filename", None)
    save_config(cfg)
    return redirect("/cv")

# ── Launch ────────────────────────────────────────────────────────────────────


@app.route("/back/<page>")
def go_back(page):
    """Generic back navigation."""
    destinations = {
        "welcome":  "/check",
        "folder":   "/welcome",
        "cv":       "/folder",
        "prefs":    "/cv",
        "connect":  "/prefs",
        "dashboard":"/connect",
    }
    return redirect(destinations.get(page, "/"))

@app.route("/instructions/view")
def view_instructions():
    """Serve instructions as plain text built from current config."""
    cfg = find_config()
    if not cfg.get("role_title"):
        return "Setup not complete. Please finish setup first.", 404
    from flask import Response
    content = build_instructions(cfg)
    return Response(content, mimetype="text/plain; charset=utf-8")


@app.route("/dashboard/reset", methods=["GET","POST"])
def reset_search():
    """Clear job preferences and job list for a fresh search."""
    if request.method == "POST":
        cfg = find_config()
        # Clear only job-search specific fields, keep folder/cv/claude settings
        for key in ["role_title","location","salary_range","work_type",
                    "instructions_path","setup_complete"]:
            cfg.pop(key, None)
        save_config(cfg)
        return redirect("/prefs")
    return render_template("reset_confirm.html")

@app.route("/prefs/saved")
def prefs_saved():
    """Confirmation page shown after updating preferences."""
    cfg = find_config()
    return render_template("prefs_saved.html", cfg=cfg)

# ── Phase 2: Job Search ───────────────────────────────────────────────────────

JSEARCH_KEY = "demo"   # placeholder - user supplies their own

def safe_filename(text):
    """Convert role title to safe filename."""
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r"[^\w\s-]", "", text).strip()
    text = re.sub(r"[\s]+", "_", text)
    return text[:50]

def get_jobs_file(cfg):
    """Return path to Excel file for the current role title."""
    app_dir  = cfg.get("app_dir", str(default_dir()))
    role     = cfg.get("role_title", "Jobs")
    filename = safe_filename(role) + "_Jobs.xlsx"
    return Path(app_dir) / "Job Listings" / filename

def get_tracker_file(cfg):
    """Return path to the single Application Tracker file."""
    app_dir = cfg.get("app_dir", str(default_dir()))
    return Path(app_dir) / "Application Tracker" / "Application_Tracker.xlsx"

def search_jsearch(role, location, salary_min, salary_max, work_type):
    """Fetch up to 100 jobs from JSearch in a single API call."""
    cfg     = find_config()
    api_key = cfg.get("jsearch_key", JSEARCH_KEY)

    loc          = (location or "").strip().lower()
    global_search = loc in ("", "global", "any", "anywhere", "worldwide")

    # Build natural language query
    if global_search:
        query = role
        if work_type and work_type.lower() == "remote":
            query = f"{role} remote"
    elif loc == "remote":
        query = f"{role} remote"
    else:
        query = f"{role} in {location}"
        if work_type and work_type.lower() == "remote":
            query += " remote"

    headers = {
        "X-RapidAPI-Key":  api_key,
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com",
    }

    # JSearch: num_pages=10 returns up to 100 results in one call
    params = {
        "query":     query,
        "page":      "1",
        "num_pages": "10",
    }

    try:
        resp = http_requests.get(
            "https://jsearch.p.rapidapi.com/search",
            headers=headers,
            params=params,
            timeout=30
        )
        resp.raise_for_status()
        data    = resp.json()
        results = data.get("data", [])

        jobs     = []
        seen_ids = set()

        for j in results:
            jid = j.get("job_id", "")
            if jid in seen_ids:
                continue
            seen_ids.add(jid)

            # Salary
            smin = j.get("job_min_salary") or 0
            smax = j.get("job_max_salary") or 0
            if smin and smax:
                salary_str = f"{int(smin):,} - {int(smax):,}"
            elif smin:
                salary_str = f"{int(smin):,}+"
            elif smax:
                salary_str = f"Up to {int(smax):,}"
            else:
                salary_str = ""

            # Client-side salary filter
            if salary_min and salary_min > 0 and smax and smax < salary_min:
                continue
            if salary_max and salary_max > 0 and smin and smin > salary_max:
                continue

            # Work type
            is_remote = j.get("job_is_remote", False)
            title_low = (j.get("job_title") or "").lower()
            desc_low  = (j.get("job_description") or "").lower()[:300]
            if is_remote or "remote" in title_low:
                wtype_str = "Remote"
            elif "hybrid" in title_low or "hybrid" in desc_low:
                wtype_str = "Hybrid"
            else:
                wtype_str = "On-site"

            # Location
            parts   = [j.get("job_city")    or "",
                       j.get("job_state")   or "",
                       j.get("job_country") or ""]
            loc_str = ", ".join(p for p in parts if p)

            # Posted date
            posted_ts = j.get("job_posted_at_datetime_utc") or ""
            posted    = posted_ts[:10] if posted_ts else ""

            jobs.append({
                "title":       j.get("job_title", ""),
                "company":     j.get("employer_name", ""),
                "location":    loc_str or location,
                "salary":      salary_str,
                "posted":      posted,
                "work_type":   wtype_str,
                "url":         (j.get("job_apply_link") or
                                j.get("job_google_link") or ""),
                "description": (j.get("job_description") or ""),
                "id":          jid,
            })

            if len(jobs) >= 100:
                break

        return jobs, len(jobs), None

    except Exception as e:
        return [], 0, str(e)


def load_existing_jobs(filepath):
    """Load existing jobs from Excel into a dict keyed by job id/url."""
    existing = {}
    if not filepath.exists() or openpyxl is None:
        return existing
    try:
        wb = openpyxl.load_workbook(str(filepath))
        ws = wb.active
        headers = [c.value for c in ws[1]]
        url_idx = headers.index("URL") if "URL" in headers else None
        id_idx  = headers.index("Job ID") if "Job ID" in headers else None
        status_idx = headers.index("Status") if "Status" in headers else None
        for row in ws.iter_rows(min_row=2, values_only=True):
            key = row[url_idx] if url_idx is not None else None
            if key:
                status = row[status_idx] if status_idx is not None else ""
                existing[key] = status or ""
    except Exception:
        pass
    return existing

def save_jobs_excel(jobs, filepath, existing_status):
    """Save jobs to Excel, newest first, max 100 rows."""
    if openpyxl is None:
        return

    filepath.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Job Listings"

    # Header: bold black font, no fill
    header_font = Font(bold=True, color="000000", size=11)
    headers = ["#", "Title", "Company", "Location", "Salary",
               "Work Type", "Posted", "Status", "URL", "Job ID"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font      = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # Column widths
    widths = [4, 35, 25, 20, 18, 12, 12, 15, 50, 15]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[
            openpyxl.utils.get_column_letter(col)].width = w

    # Write jobs - black font, no fill
    plain_font = Font(color="000000", size=10)
    for i, job in enumerate(jobs[:100], 1):
        status   = existing_status.get(job["url"], "")
        row_data = [i, job["title"], job["company"], job["location"],
                    job["salary"], job["work_type"], job["posted"],
                    status, job["url"], job["id"]]
        for col, val in enumerate(row_data, 1):
            cell           = ws.cell(row=i+1, column=col, value=val)
            cell.font      = plain_font
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.freeze_panes = "A2"
    wb.save(str(filepath))

def ensure_tracker(cfg):
    """Create Application Tracker Excel if it doesn't exist."""
    if openpyxl is None:
        return
    filepath = get_tracker_file(cfg)
    if filepath.exists():
        return
    filepath.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Applications"

    header_font = Font(bold=True, color="000000", size=11)
    headers = ["Date Applied", "Company", "Role", "Location",
               "Work Type", "Salary", "Application URL",
               "Status", "Notes"]
    widths  = [14, 25, 30, 20, 12, 18, 50, 15, 40]

    for col, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font  = header_font
        cell.alignment = Alignment(horizontal="center")
        ws.column_dimensions[
            openpyxl.utils.get_column_letter(col)].width = w

    ws.freeze_panes = "A2"
    wb.save(str(filepath))

def parse_salary(salary_str):
    """Parse salary range string into (min, max) integers."""
    if not salary_str or salary_str.lower() == "not specified":
        return 0, 0
    nums = re.findall(r"[\d]+", salary_str.replace(",", ""))
    if len(nums) >= 2:
        return int(nums[0]), int(nums[1])
    elif len(nums) == 1:
        return int(nums[0]), 0
    return 0, 0


def build_claude_prompt(title, company, desc, cfg, cv_text=""):
    """Build a strict single-response Claude prompt for CV tailoring."""
    role     = cfg.get("role_title",   title)
    location = cfg.get("location",     "")
    salary   = cfg.get("salary_range", "")

    cv_section = ""
    if cv_text and len(cv_text.strip()) > 50:
        cv_section = f"""
MY CURRENT CV:
---
{cv_text[:3000]}
---
"""
    else:
        cv_section = """
MY CV: [Attached to this Claude Project as a file]
"""

    prompt = f"""You are an expert CV writer and career coach. I need you to tailor my CV and write a cover letter for a specific role. Do this immediately with what I give you — do not ask follow-up questions.

ROLE: {title}
COMPANY: {company}
LOCATION: {location}
SALARY RANGE: {salary}

JOB DESCRIPTION:
---
{desc[:2000] if desc else "See job details above"}
---
{cv_section}
INSTRUCTIONS:
1. HONEST FIT ASSESSMENT (do this first):
   - Review my CV against the job requirements carefully
   - Give a clear honest fit score out of 10
   - If fit is below 5, say: "Your profile is not a strong match for this role.
     Gaps: [list specific gaps]. You may want to reconsider applying."
   - If fit is 5 or above, say: "You are a [X]/10 match. Strengths: [list].
     Gaps: [list any]."
   - Be specific and honest - name exact requirements met and not met

2. Extract the top 10 ATS keywords from the job description
3. Tailor my CV to match this role - preserve all facts, never invent experience
4. Score the tailored CV against the job (target 80+/100), improve if below 80
5. Write a fresh specific cover letter (maximum 1 page)
6. In the cover letter, address any gaps honestly - do not oversell

CRITICAL OUTPUT RULES:
- Output ONLY plain text in this chat — do NOT create, attach, or offer any file downloads
- Do NOT use Canvas, artifacts, or document editors
- Write everything directly in your chat response as plain text
- The user will copy and paste from your response

OUTPUT FORMAT — use exactly these section headers with no other text before or after:

TAILORED CV:
[write the complete tailored CV here as plain text]

COVER LETTER:
[write the complete cover letter here as plain text]

END OF DOCUMENTS

---
**YOUR DOCUMENTS ARE READY — NEXT STEPS:**
1. Copy all the text under "TAILORED CV:" and paste it into the Tailored CV box in AJAS
2. Copy all the text under "COVER LETTER:" and paste it into the Cover Letter box in AJAS
3. Click "Save as Word Documents" in AJAS — files will be saved automatically to:
   - Tailored CV → My CV > Tailored CVs folder
   - Cover Letter → Cover Letters folder
4. Review and edit the saved files if needed, then return to AJAS to open the application
---"""

    return prompt


def _search_cache_path():
    return Path.home() / ".ajas_last_search.json"

def _save_last_search(jobs, params):
    """Save search results to temp file for back-navigation."""
    try:
        data = {"jobs": jobs, "params": params}
        with open(_search_cache_path(), "w") as f:
            json.dump(data, f)
    except Exception:
        pass

def _load_last_search():
    """Load last search results from temp file."""
    try:
        p = _search_cache_path()
        if p.exists():
            with open(p) as f:
                data = json.load(f)
            return data.get("jobs", []), data.get("params", {})
    except Exception:
        pass
    return [], {}


# ── Phase 2 Routes ────────────────────────────────────────────────────────────

@app.route("/search")
def search_page():
    """Job search page - checks for API key first."""
    cfg = find_config()
    if not cfg.get("jsearch_key"):
        return redirect("/search/setup")
    return render_template("search.html", cfg=cfg, jobs=None, error=None,
                           total=0, searched=False)

@app.route("/search/setup", methods=["GET", "POST"])
def search_setup():
    """First-time JSearch API key setup."""
    if request.method == "POST":
        app_id  = request.form.get("app_id",  "").strip()
        app_key = request.form.get("app_key", "").strip()
        if not app_id:
            return render_template("search_setup.html",
                error="Please paste your RapidAPI key.")

        # Save key and test it
        cfg = find_config()
        cfg["jsearch_key"] = app_id   # app_id field reused for the single key
        save_config(cfg)

        try:
            headers = {
                "X-RapidAPI-Key":  app_id,
                "X-RapidAPI-Host": "jsearch.p.rapidapi.com",
            }
            resp = http_requests.get(
                "https://jsearch.p.rapidapi.com/search",
                headers=headers,
                params={"query": "developer", "num_pages": "1"},
                timeout=10
            )
            if resp.status_code == 401 or resp.status_code == 403:
                cfg.pop("jsearch_key", None)
                save_config(cfg)
                return render_template("search_setup.html",
                    error="API key not accepted. Please check your RapidAPI key "
                          "and make sure you have subscribed to the JSearch API.")
            return redirect("/search")
        except Exception:
            return redirect("/search")

    cfg = find_config()
    return render_template("search_setup.html", error=None,
        saved_id=cfg.get("jsearch_key", ""))

@app.route("/search/run", methods=["POST"])
def search_run():
    """Execute job search and return results."""
    cfg      = find_config() or {}
    role     = request.form.get("role",     "").strip() or cfg.get("role_title",   "")
    location = request.form.get("location", "").strip() or cfg.get("location",     "")
    salary   = request.form.get("salary",   "").strip() or cfg.get("salary_range", "")
    wtype    = request.form.get("work_type","").strip() or cfg.get("work_type",    "Any")
    if not role:
        return render_template("search.html", cfg=cfg,
            jobs=None, error="Please set a role title in Job Preferences first.",
            total=0, searched=True, jobs_file="",
            role="", location="", salary="", wtype="Any")

    sal_min, sal_max = parse_salary(salary)
    jobs, total, err = search_jsearch(role, location, sal_min, sal_max, wtype)
    # Store results in a temp file for back-navigation
    if jobs:
        _save_last_search(jobs, {
            "role": role, "location": location,
            "salary": salary, "wtype": wtype,
            "jobs_file": str(get_jobs_file(cfg))
        })

    if err:
        if "401" in str(err) or "403" in str(err):
            friendly = ("API key error — please go back to the dashboard, "
                        "click Search for Jobs, and re-enter your RapidAPI key.")
        elif "timeout" in str(err).lower() or "connect" in str(err).lower():
            friendly = ("Connection timed out — please check your internet "
                        "connection and try again.")
        else:
            friendly = ("Search could not complete. Please try again. "
                        f"Details: {err}")
        return render_template("search.html", cfg=cfg,
            jobs=None, error=friendly, total=0, searched=True,
            role=role, location=location, salary=salary, wtype=wtype)

    # Load existing file to preserve applied status
    jobs_file = get_jobs_file(cfg)
    existing  = load_existing_jobs(jobs_file)

    # Mark previously applied jobs
    for job in jobs:
        job["status"] = existing.get(job["url"], "")

    # Save to Excel
    save_jobs_excel(jobs, jobs_file, existing)

    # Ensure tracker exists
    ensure_tracker(cfg)

    return render_template("search.html", cfg=cfg,
        jobs=jobs[:100], error=None, total=total,
        searched=True, jobs_file=str(jobs_file),
        role=role, location=location, salary=salary, wtype=wtype)

@app.route("/apply/<path:job_url>")
def apply_job(job_url):
    """Show apply page with Claude prompt and paste areas."""
    from urllib.parse import unquote
    cfg     = find_config()
    job_url = unquote(job_url)
    title   = request.args.get("title",   "")
    company = request.args.get("company", "")
    desc    = request.args.get("desc",    "")

    cv_filename = cfg.get("cv_filename", "")
    cv_path     = cfg.get("cv_path",     "")

    # Read master CV text if available
    cv_text = ""
    if cv_path:
        try:
            p = Path(cv_path)
            if p.suffix.lower() == ".docx":
                from docx import Document
                doc = Document(str(p))
                cv_text = "\n".join(para.text for para in doc.paragraphs
                                     if para.text.strip())
            elif p.suffix.lower() == ".pdf":
                cv_text = "[PDF CV - see filename: " + p.name + "]"
        except Exception:
            cv_text = ""

    # Build the strict single-response Claude prompt
    claude_prompt = build_claude_prompt(
        title, company, desc, cfg, cv_text
    )

    app_dir      = Path(cfg.get("app_dir", str(default_dir())))
    cv_save_path = str(app_dir / "My CV" / "Tailored CVs" /
        (safe_filename(company) + "_" + safe_filename(title) + "_CV.docx"))
    cl_save_path = str(app_dir / "Cover Letters" /
        (safe_filename(company) + "_" + safe_filename(title) + "_Cover_Letter.docx"))
    return render_template("apply.html", cfg=cfg,
        job_url=job_url, title=title, company=company,
        desc=desc, claude_prompt=claude_prompt,
        cv_save_path=cv_save_path, cl_save_path=cl_save_path)

@app.route("/apply/open", methods=["POST"])
def apply_open():
    """Mark job as applied, update Excel, open URL in new tab via JS."""
    from urllib.parse import unquote
    from flask import Response
    job_url = unquote(request.form.get("job_url", ""))
    company = request.form.get("company", "")
    title   = request.form.get("title",   "")

    # Update Excel status
    if job_url:
        cfg = find_config()
        jobs_file = get_jobs_file(cfg)
        if jobs_file.exists() and openpyxl is not None:
            try:
                wb = openpyxl.load_workbook(str(jobs_file))
                ws = wb.active
                headers    = [c.value for c in ws[1]]
                url_idx    = headers.index("URL")    + 1 if "URL"    in headers else None
                status_idx = headers.index("Status") + 1 if "Status" in headers else None
                if url_idx and status_idx:
                    for row in ws.iter_rows(min_row=2):
                        if row[url_idx-1].value == job_url:
                            row[status_idx-1].value = "Applied"
                            break
                wb.save(str(jobs_file))
            except Exception:
                pass

    # Store session for confirm page
    from urllib.parse import quote as url_quote
    from flask import session
    session["confirm_title"]   = title
    session["confirm_company"] = company
    session["confirm_job_url"] = job_url
    session["confirm_fresh"]   = True

    safe_title   = url_quote(title,   safe="")
    safe_company = url_quote(company, safe="")
    safe_job_url = url_quote(job_url, safe="")

    # Redirect straight to confirm page - no intermediate page, no extra tabs.
    # The job URL was already opened by the form's target="_blank" in apply.html
    return redirect(f"/apply/confirm?title={safe_title}"
                    f"&company={safe_company}&job_url={safe_job_url}")

@app.route("/search/results")
def search_results():
    """Restore last search results from temp file."""
    cfg    = find_config()
    jobs, params = _load_last_search()
    return render_template("search.html", cfg=cfg,
        jobs=jobs, error=None,
        total=len(jobs), searched=bool(jobs),
        jobs_file=params.get("jobs_file", str(get_jobs_file(cfg))),
        role=params.get("role",""),
        location=params.get("location",""),
        salary=params.get("salary",""),
        wtype=params.get("wtype","Any"))


@app.route("/search/open_file")
def search_open_file():
    """Open the job listings Excel file and stay on results."""
    cfg = find_config()
    jobs_file = get_jobs_file(cfg)
    if jobs_file.exists():
        open_in_finder(jobs_file)
    return redirect("/search/results")


@app.route("/apply/save_docs", methods=["POST"])
def apply_save_docs():
    """Save pasted tailored CV and cover letter as .docx files."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cfg     = find_config()
    app_dir = Path(cfg.get("app_dir", str(default_dir())))
    title   = request.form.get("title",   "Role").strip()
    company = request.form.get("company", "Company").strip()
    cv_text = request.form.get("cv_text", "").strip()
    cl_text = request.form.get("cl_text", "").strip()

    if not cv_text and not cl_text:
        return jsonify({"error": "Please paste at least one document."})

    saved = []
    errors = []

    def get_original_cv_style():
        """Extract font name and base size from the original CV if available."""
        cv_path = cfg.get("cv_path", "")
        font_name = "Calibri"
        base_size = Pt(11)
        try:
            if cv_path and Path(cv_path).suffix.lower() == ".docx":
                orig = Document(cv_path)
                # Sample first few runs to get the dominant font
                for para in orig.paragraphs[:20]:
                    for run in para.runs:
                        if run.font.name and run.font.name.strip():
                            font_name = run.font.name
                        if run.font.size:
                            base_size = run.font.size
                        break
                    else:
                        continue
                    break
        except Exception:
            pass
        return font_name, base_size

    def make_docx(text, filepath):
        """Create a .docx that matches the original CV font and style."""
        font_name, base_size = get_original_cv_style()

        doc = Document()
        # Match original CV margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # Set default font to match original
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        style = doc.styles["Normal"]
        style.font.name      = font_name
        style.font.size      = base_size
        style.font.color.rgb = None   # black

        for line in text.split("\n"):
            line = line.rstrip()
            stripped = line.strip()
            p = doc.add_paragraph()

            # Section headings: ALL CAPS or short line ending with colon
            if stripped and (stripped.isupper() or
               (stripped.endswith(":") and len(stripped) < 60
                and not stripped.startswith("-"))):
                run = p.add_run(stripped)
                run.bold            = True
                run.font.name       = font_name
                run.font.size       = Pt(base_size.pt + 1) if base_size else Pt(12)
                run.font.color.rgb  = None

            # Bullet points
            elif stripped.startswith(("- ", "• ", "* ")):
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run("•  " + stripped[2:])
                run.font.name      = font_name
                run.font.size      = base_size
                run.font.color.rgb = None

            # Bold inline (text wrapped in ** **)
            elif "**" in stripped:
                parts = stripped.split("**")
                for idx2, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.bold            = (idx2 % 2 == 1)
                        run.font.name       = font_name
                        run.font.size       = base_size
                        run.font.color.rgb  = None

            # Normal text
            else:
                run = p.add_run(line)
                run.font.name      = font_name
                run.font.size      = base_size
                run.font.color.rgb = None

            p.paragraph_format.space_after = Pt(2)

        doc.save(str(filepath))

    # Save tailored CV
    if cv_text:
        cv_dir = app_dir / "My CV" / "Tailored CVs"
        cv_dir.mkdir(parents=True, exist_ok=True)
        fname   = safe_filename(company) + "_" + safe_filename(title) + "_CV.docx"
        cv_path = cv_dir / fname
        try:
            make_docx(cv_text, cv_path)
            saved.append(("Tailored CV", str(cv_path)))
        except Exception as e:
            errors.append(f"CV: {e}")

    # Save cover letter
    if cl_text:
        cl_dir = app_dir / "Cover Letters"
        cl_dir.mkdir(parents=True, exist_ok=True)
        fname   = safe_filename(company) + "_" + safe_filename(title) + "_Cover_Letter.docx"
        cl_path = cl_dir / fname
        try:
            make_docx(cl_text, cl_path)
            saved.append(("Cover Letter", str(cl_path)))
        except Exception as e:
            errors.append(f"Cover Letter: {e}")

    if errors and not saved:
        return jsonify({"error": " | ".join(errors)})

    response = {"saved": saved}
    if errors:
        response["warnings"] = errors
    return jsonify(response)

@app.route("/apply/open_tailored_folder")
def open_tailored_folder():
    """Open the Tailored CVs folder."""
    cfg     = find_config()
    app_dir = Path(cfg.get("app_dir", str(default_dir())))
    folder  = app_dir / "My CV" / "Tailored CVs"
    folder.mkdir(parents=True, exist_ok=True)
    open_in_finder(folder)
    return ("", 204)


@app.route("/apply/confirm")
def apply_confirm():
    """Post-application confirmation and tracker update."""
    from flask import session
    cfg     = find_config()
    title   = request.args.get("title",   "")
    company = request.args.get("company", "")
    job_url = request.args.get("job_url", "")
    logged  = request.args.get("logged",  "")
    today   = str(date.today())
    # Only show "Application Opened" banner on fresh arrival
    fresh   = session.pop("confirm_fresh", False)
    return render_template("apply_confirm.html",
        cfg=cfg, title=title, company=company,
        job_url=job_url, logged=logged, fresh=fresh,
        today=today, tracker_file=str(get_tracker_file(cfg)))

@app.route("/apply/log", methods=["POST"])
def apply_log():
    """Log application to the tracker Excel file."""
    from datetime import date as dt_date
    cfg     = find_config()
    title   = request.form.get("title",   "")
    company = request.form.get("company", "")
    notes   = request.form.get("notes",   "")

    # Ensure tracker exists
    ensure_tracker(cfg)
    tracker = get_tracker_file(cfg)

    if openpyxl is not None and tracker.exists():
        try:
            wb = openpyxl.load_workbook(str(tracker))
            ws = wb.active
            ws.append([
                str(dt_date.today()),          # Date Applied
                company,                        # Company
                title,                          # Role
                cfg.get("location",     ""),   # Location
                cfg.get("work_type",    ""),   # Work Type
                cfg.get("salary_range", ""),   # Salary
                request.form.get("job_url",""),# URL
                "Applied",                      # Status
                notes,                          # Notes
            ])
            wb.save(str(tracker))
        except Exception:
            pass

    # Also mark as Applied in the job listings Excel
    job_url_val = request.form.get("job_url", "")
    if job_url_val:
        jobs_file = get_jobs_file(cfg)
        if jobs_file.exists() and openpyxl is not None:
            try:
                wb = openpyxl.load_workbook(str(jobs_file))
                ws = wb.active
                headers    = [c.value for c in ws[1]]
                url_idx    = headers.index("URL")    + 1 if "URL"    in headers else None
                status_idx = headers.index("Status") + 1 if "Status" in headers else None
                if url_idx and status_idx:
                    for row in ws.iter_rows(min_row=2):
                        if row[url_idx-1].value == job_url_val:
                            row[status_idx-1].value = "Applied"
                            break
                wb.save(str(jobs_file))
            except Exception:
                pass
        # Update the search cache too so results page shows Applied
        jobs, params = _load_last_search()
        for j in jobs:
            if j.get("url") == job_url_val:
                j["status"] = "Applied"
                break
        _save_last_search(jobs, params)

    from urllib.parse import quote
    return redirect("/apply/confirm?title=" + quote(title, safe="") +
                    "&company=" + quote(company, safe="") +
                    "&job_url=" + quote(job_url_val, safe="") +
                    "&logged=1")

@app.route("/apply/open_tracker")
def open_tracker():
    """Open the Application Tracker Excel file."""
    cfg     = find_config()
    tracker = get_tracker_file(cfg)
    ensure_tracker(cfg)
    open_in_finder(tracker)
    return ("", 204)


@app.route("/tracker")
def tracker_page():
    """Application tracker management page — reads ALL tracker files."""
    cfg     = find_config()
    app_dir = Path(cfg.get("app_dir", str(default_dir())))
    tracker_dir = app_dir / "Application Tracker"
    tracker_dir.mkdir(parents=True, exist_ok=True)

    entries = []

    if openpyxl is not None:
        tracker_file_path = get_tracker_file(cfg)
        ensure_tracker(cfg)
        if tracker_file_path.exists():
            try:
                wb      = openpyxl.load_workbook(str(tracker_file_path))
                ws      = wb.active
                headers = [c.value for c in ws[1]]
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if any(cell for cell in row):
                        entry = dict(zip(headers, row))
                        entries.append(entry)
            except Exception:
                pass

    # Sort: most recent date first
    def sort_key(e):
        d = str(e.get("Date Applied") or "")
        return d
    entries.sort(key=sort_key, reverse=True)

    today        = str(date.today())
    tracker_file = str(get_tracker_file(cfg))
    return render_template("tracker.html", cfg=cfg,
        entries=entries, tracker_file=tracker_file, today=today)

@app.route("/tracker/update", methods=["POST"])
def tracker_update():
    """Update the status and date of an application in the tracker."""
    cfg        = find_config()
    company    = request.form.get("company",    "").strip()
    title      = request.form.get("title",      "").strip()
    new_status = request.form.get("status",     "").strip()
    status_date= request.form.get("status_date","").strip()
    notes      = request.form.get("notes",      "").strip()
    source     = request.form.get("source",     "confirm")

    app_dir     = Path(cfg.get("app_dir", str(default_dir())))
    tracker_dir = app_dir / "Application Tracker"
    tracker_dir.mkdir(parents=True, exist_ok=True)
    updated     = False

    if openpyxl is not None:
        tfile = get_tracker_file(cfg)
        ensure_tracker(cfg)
        if tfile.exists():
            try:
                wb      = openpyxl.load_workbook(str(tfile))
                ws      = wb.active
                headers = [c.value for c in ws[1]]
                col     = {h: i for i, h in enumerate(headers)}

                for row in ws.iter_rows(min_row=2):
                    row_company = str(row[col["Company"]].value or "").strip() if "Company" in col else ""
                    row_title   = str(row[col["Role"]].value    or "").strip() if "Role"    in col else ""
                    if row_company == company and row_title == title:
                        if "Status" in col:
                            row[col["Status"]].value = new_status
                        if "Notes" in col:
                            existing   = str(row[col["Notes"]].value or "").strip()
                            note_entry = f"{status_date}: {new_status}"
                            if notes:
                                note_entry += f" ({notes})"
                            row[col["Notes"]].value = (
                                f"{existing} | {note_entry}"
                                if existing else note_entry
                            )
                        updated = True
                        break
                if updated:
                    wb.save(str(tfile))
            except Exception:
                pass

    if source == "tracker":
        return redirect("/tracker?updated=" + company)

    # Return to confirmation page
    from urllib.parse import quote
    return redirect("/apply/confirm?title=" + quote(title, safe="") +
                    "&company=" + quote(company, safe="") +
                    "&logged=1&status_updated=1&new_status=" +
                    quote(new_status, safe=""))


@app.route("/settings", methods=["GET", "POST"])
def settings():
    """Settings page - update API keys and preferences."""
    cfg = find_config()
    msg = None
    if request.method == "POST":
        action = request.form.get("action", "")
        if action == "update_api":
            key        = request.form.get("jsearch_key",  "").strip()
            claude_url = request.form.get("claude_url",   "").strip()
            changed    = False
            if key:
                cfg["jsearch_key"] = key
                changed = True
            if claude_url and claude_url.startswith("http"):
                cfg["claude_project_url"] = claude_url
                changed = True
            if changed:
                save_config(cfg)
                msg = "Settings updated successfully."
            else:
                msg = "No changes to save — please enter at least one value."
        elif action == "clear_api":
            cfg.pop("jsearch_key", None)
            save_config(cfg)
            msg = "API key cleared. You will be prompted to enter it next time you search."
    return render_template("settings.html", cfg=cfg, msg=msg)

@app.route("/dashboard/settings")
def dashboard_settings():
    return redirect("/settings")


def open_browser():
    time.sleep(1.2)
    try:
        webbrowser.open(f"http://127.0.0.1:{PORT}")
    except Exception:
        pass

if __name__ == "__main__":
    import logging
    import socket
    logging.getLogger("werkzeug").setLevel(logging.ERROR)

    # Always work from the script directory
    script_dir = Path(__file__).parent.resolve()
    os.chdir(str(script_dir))
    app.template_folder = str(script_dir / "templates")
    app.static_folder   = str(script_dir / "static")

    # Find a free port - simple and fast, no blocking
    def find_free_port(preferred):
        for port in [preferred, preferred+1, preferred+2,
                     preferred+10, preferred+100, 8000, 8080, 8888]:
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
                s.bind(("127.0.0.1", port))
                s.close()
                return port
            except OSError:
                continue
        return preferred + 200  # last resort

    actual_port = find_free_port(PORT)

    # Open browser after short delay
    def open_browser_delayed():
        time.sleep(1.5)
        url = f"http://127.0.0.1:{actual_port}"
        try:
            # Most reliable cross-platform approach
            webbrowser.open(url)
        except Exception:
            pass

    threading.Thread(target=open_browser_delayed, daemon=True).start()
    app.run(host="127.0.0.1", port=actual_port,
            debug=False, use_reloader=False)
